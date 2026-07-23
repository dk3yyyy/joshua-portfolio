import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Mm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Flowable, KeepTogether, PageBreak, Paragraph, SimpleDocTemplate

ROOT = Path(__file__).resolve().parents[1]
PUBLIC = ROOT / "public"
PUBLIC.mkdir(parents=True, exist_ok=True)
DOCX_PATH = PUBLIC / "Joshua_Nwachinemere_CV.docx"
PDF_PATH = PUBLIC / "Joshua_Nwachinemere_CV.pdf"

NAME = "JOSHUA NWACHINEMERE"
HEADLINE = "AI Engineer | Applied AI, ML Systems, Automation"
EMAIL = "josh0victor@outlook.com"
PHONE = "+234 913 148 0096"
LOCATION = "Lagos, Nigeria | Open to UK relocation"
GITHUB_URL = "https://github.com/dk3yyyy"
LINKEDIN_URL = "https://www.linkedin.com/in/joshua-nwachinemere"
PORTFOLIO_URL = "https://dk3yyyy.github.io/joshua-portfolio"
URL_PATTERN = re.compile(r"https://[^\s|]+")

SUMMARY = (
    "AI Engineer and VolyxAI founder building applied AI, backend services, and automation systems with Python, "
    "JavaScript, and TypeScript. Experience includes a multimodal macOS assistant, retrieval workflows, "
    "real-time applications, ML evaluation, API integrations, and human-in-the-loop automation."
)

EXPERIENCE = [
    {
        "role": "Founder & AI Automation Engineer",
        "org": "VolyxAI",
        "date": "Nov 2025 - Present",
        "bullets": [
            "Built and tested controlled workflow prototypes connecting conversational intake to schema validation, calendar scheduling, email confirmation, and approval-based handoffs.",
            "Integrated language models, n8n, webhooks, Vapi, ElevenLabs, Google Calendar, and email services in controlled test environments.",
            "Designed deterministic validation, bounded retries, idempotency controls, and human approval for consequential actions.",
        ],
        "tools": "Python, n8n, TypeScript, REST APIs, webhooks, Vapi, ElevenLabs, Azure AI Foundry",
    },
    {
        "role": "Independent Software & Automation Developer",
        "org": "Self-Employed / Independent Projects",
        "date": "Jan 2021 - Present",
        "bullets": [
            "Built Python automation and backend utilities for data processing, Telegram bots, API integrations, and operational tooling.",
            "Developed stream-oriented utilities for processing large text datasets without loading complete files into memory.",
            "Created asynchronous integrations for wallet analysis, malware scanning, user tracking, pagination, caching, and admin reporting.",
        ],
        "tools": "Python, Async I/O, Telegram Bot API, aiohttp, SQL, Docker, REST APIs",
    },
]

TRAINING = [
    {
        "name": "Threat Detection & Response Capstone | Capstone IT Staffing (Remote)",
        "date": "Jul 2025 - Aug 2025",
        "bullets": [
            "Deployed Suricata on Ubuntu, wrote custom Snort rules, and automated alert enrichment with Python and public OSINT sources.",
            "Built a DNS intelligence CLI with asynchronous resolution, WHOIS lookups, trace output, and JSON export.",
        ],
    }
]

PROJECTS = [
    {
        "name": "Volyx Lens | Privacy-First Context-Aware AI Assistant",
        "stack": "JavaScript, Electron, Swift, Azure AI Foundry, Multimodal APIs",
        "bullets": [
            "Developed and substantially extended a macOS assistant that combines explicitly selected screen context, microphone input, and meeting audio with user-configured AI providers.",
            "Implemented local OCR, bounded screenshot memory, context ranking, provider routing, and consent before larger image uploads.",
            "Hardened Electron boundaries with sandboxing, context isolation, restricted IPC, automated tests, secret scanning, and release checks.",
            "Source: https://github.com/dk3yyyy/volyx-lens",
        ],
    },
    {
        "name": "Football Predictor | Experimental ML Pipeline",
        "stack": "Python, XGBoost, scikit-learn, FastAPI, Streamlit, SQLAlchemy",
        "bullets": [
            "Built data scrapers, temporal feature engineering, XGBoost outcome and Poisson goal models, a FastAPI service, and a Streamlit dashboard.",
            "Evaluated with rolling-origin testing across 1,140 Premier League matches: 53.77% outcome accuracy versus a 56.70% bookmaker-implied benchmark, showing signal but no practical betting advantage.",
            "Archived evidence: https://github.com/dk3yyyy/football_predictor/tree/repair/football-predictor-hardening",
        ],
    },
    {
        "name": "Noughtline | Real-Time Multiplayer System",
        "stack": "React, Express, Socket.IO, SQLite, Paystack",
        "bullets": [
            "Built signed guest sessions, server-authoritative moves, reconnect and forfeit handling, transactional match settlement, and an append-only currency ledger.",
            "Added idempotent Paystack crediting plus automated API, economy, room-state, and two-client Socket.IO tests.",
            "Live: https://noughtline.onrender.com | Source: https://github.com/dk3yyyy/Noughtline",
        ],
    },
    {
        "name": "VirusTotal Telegram Bot | Security API Integration",
        "stack": "Python, Pyrogram, aiohttp, VirusTotal API",
        "bullets": [
            "Built asynchronous file, URL, and hash scanning with local SHA-256 lookup, VirusTotal large-file upload handling, bounded retry backoff, and short-lived report caching.",
            "Source: https://github.com/dk3yyyy/VirusTotal-Telegram-Bot",
        ],
    },
]

SKILLS = [
    "Languages: Python, JavaScript, TypeScript, SQL",
    "AI / ML: Azure AI Foundry, XGBoost, scikit-learn, pandas, Ollama, LangChain, Chroma, RAG, model evaluation",
    "Backend & Automation: FastAPI, Express, Socket.IO, n8n, REST APIs, webhooks, SQLAlchemy, SQLite",
    "Infrastructure: Docker, GitHub Actions, Linux, Microsoft Azure, Electron",
    "Security: Suricata, Snort, VirusTotal API, OSINT, DNS analysis, secure credential management",
]

CERTIFICATIONS = [
    "Google AI Professional Certificate | Coursera",
    "Google Cybersecurity Professional Certificate | Coursera",
]

EDUCATION = [
    "Bachelor of Technology (BTech), Mathematics | 2016 - 2021",
    "Federal University of Technology, Owerri (FUTO)",
]


def configure_docx() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    section.different_first_page_header_footer = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_run("Joshua Nwachinemere | AI Engineer").font.size = Pt(8)
    doc.core_properties.title = "Joshua Nwachinemere CV"
    doc.core_properties.author = "Joshua Nwachinemere"
    doc.core_properties.subject = "AI Engineer CV"
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.2)
    return doc


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_text_with_links(paragraph, text: str) -> None:
    cursor = 0
    for match in URL_PATTERN.finditer(text):
        paragraph.add_run(text[cursor:match.start()])
        url = match.group(0)
        add_hyperlink(paragraph, url, url)
        cursor = match.end()
    paragraph.add_run(text[cursor:])


def add_docx_contact(doc: Document) -> None:
    first = doc.add_paragraph()
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first.paragraph_format.space_after = Pt(0)
    first.add_run(f"{LOCATION} | ")
    add_hyperlink(first, EMAIL, f"mailto:{EMAIL}")
    first.add_run(f" | {PHONE}")

    second = doc.add_paragraph()
    second.alignment = WD_ALIGN_PARAGRAPH.CENTER
    second.paragraph_format.space_after = Pt(4)
    add_hyperlink(second, "github.com/dk3yyyy", GITHUB_URL)
    second.add_run(" | ")
    add_hyperlink(second, "linkedin.com/in/joshua-nwachinemere", LINKEDIN_URL)
    second.add_run(" | ")
    add_hyperlink(second, "dk3yyyy.github.io/joshua-portfolio", PORTFOLIO_URL)


def add_docx_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)


def add_docx_bullets(doc: Document, bullets: list[str]) -> None:
    for bullet in bullets:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(0)
        add_text_with_links(p, f"- {bullet}")


def build_docx() -> None:
    doc = configure_docx()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(NAME)
    r.bold = True
    r.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(HEADLINE)
    r.bold = True
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(1)

    add_docx_contact(doc)

    add_docx_heading(doc, "PROFILE")
    p = doc.add_paragraph(SUMMARY)
    p.paragraph_format.space_after = Pt(2)

    add_docx_heading(doc, "CORE SKILLS")
    add_docx_bullets(doc, SKILLS)

    add_docx_heading(doc, "PROFESSIONAL EXPERIENCE")
    for item in EXPERIENCE:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(item["role"])
        r.bold = True
        p.add_run(f" | {item['date']}").italic = True
        p = doc.add_paragraph(item["org"])
        p.paragraph_format.space_after = Pt(0)
        add_docx_bullets(doc, item["bullets"])
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("Key tools: ")
        r.bold = True
        p.add_run(item["tools"])

    add_docx_heading(doc, "TECHNICAL TRAINING")
    for item in TRAINING:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(item["name"])
        r.bold = True
        p.add_run(f" | {item['date']}").italic = True
        add_docx_bullets(doc, item["bullets"])

    doc.add_page_break()
    add_docx_heading(doc, "SELECTED PROJECTS")
    for project in PROJECTS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(project["name"])
        r.bold = True
        p = doc.add_paragraph(project["stack"])
        p.paragraph_format.space_after = Pt(0)
        p.runs[0].italic = True
        add_docx_bullets(doc, project["bullets"])

    add_docx_heading(doc, "CERTIFICATIONS")
    add_docx_bullets(doc, CERTIFICATIONS)
    add_docx_heading(doc, "EDUCATION")
    for line in EDUCATION:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)

    doc.save(DOCX_PATH)


def pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "name": ParagraphStyle("CVName", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=16, leading=18, alignment=TA_CENTER, spaceAfter=2),
        "headline": ParagraphStyle("CVHeadline", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=10.5, leading=12, alignment=TA_CENTER, spaceAfter=2),
        "contact": ParagraphStyle("CVContact", parent=styles["Normal"], fontName="Helvetica", fontSize=8.1, leading=10, alignment=TA_CENTER, spaceAfter=5),
        "section": ParagraphStyle("CVSection", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=colors.HexColor("#111111"), spaceBefore=10, spaceAfter=5, borderWidth=0, borderPadding=0),
        "role": ParagraphStyle("CVRole", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=10.8, leading=13.5, spaceBefore=5, spaceAfter=1.5),
        "org": ParagraphStyle("CVOrg", parent=styles["Normal"], fontName="Helvetica-Oblique", fontSize=10, leading=12.5, spaceAfter=1.5),
        "body": ParagraphStyle("CVBody", parent=styles["Normal"], fontName="Helvetica", fontSize=10.7, leading=13.7, spaceAfter=2.5),
        "tools": ParagraphStyle("CVTools", parent=styles["Normal"], fontName="Helvetica", fontSize=9.9, leading=12.5, spaceAfter=5),
    }


def bullet_list(lines: list[str], style: ParagraphStyle) -> KeepTogether:
    bullet_style = ParagraphStyle(
        f"{style.name}Bullet",
        parent=style,
        leftIndent=10,
        firstLineIndent=-7,
        spaceAfter=1,
    )
    return KeepTogether(
        [
            Paragraph(
                "- "
                + URL_PATTERN.sub(
                    lambda match: f'<link href="{match.group(0)}" color="#1155CC">{match.group(0)}</link>',
                    line,
                ),
                bullet_style,
            )
            for line in lines
        ]
    )


def draw_later_page_header(canvas, document) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#444444"))
    canvas.drawRightString(A4[0] - 14 * mm, A4[1] - 8 * mm, "Joshua Nwachinemere | AI Engineer")
    canvas.setStrokeColor(colors.HexColor("#CCCCCC"))
    canvas.line(14 * mm, A4[1] - 10 * mm, A4[0] - 14 * mm, A4[1] - 10 * mm)
    canvas.restoreState()


def build_pdf() -> None:
    styles = pdf_styles()
    contact_markup = (
        f'{LOCATION} | <link href="mailto:{EMAIL}" color="#1155CC">{EMAIL}</link> | {PHONE}<br/>'
        f'<link href="{GITHUB_URL}" color="#1155CC">github.com/dk3yyyy</link> | '
        f'<link href="{LINKEDIN_URL}" color="#1155CC">linkedin.com/in/joshua-nwachinemere</link> | '
        f'<link href="{PORTFOLIO_URL}" color="#1155CC">dk3yyyy.github.io/joshua-portfolio</link>'
    )
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=14 * mm,
        rightMargin=14 * mm,
        topMargin=14 * mm,
        bottomMargin=10 * mm,
        title="Joshua Nwachinemere CV",
        author="Joshua Nwachinemere",
    )
    story: list[Flowable] = [
        Paragraph(NAME, styles["name"]),
        Paragraph(HEADLINE, styles["headline"]),
        Paragraph(contact_markup, styles["contact"]),
        Paragraph("PROFILE", styles["section"]),
        Paragraph(SUMMARY, styles["body"]),
        Paragraph("CORE SKILLS", styles["section"]),
        bullet_list(SKILLS, styles["body"]),
        Paragraph("PROFESSIONAL EXPERIENCE", styles["section"]),
    ]
    for item in EXPERIENCE:
        story.append(Paragraph(f"{item['role']} | {item['date']}", styles["role"]))
        story.append(Paragraph(item["org"], styles["org"]))
        story.append(bullet_list(item["bullets"], styles["body"]))
        story.append(Paragraph(f"<b>Key tools:</b> {item['tools']}", styles["tools"]))

    story.append(Paragraph("TECHNICAL TRAINING", styles["section"]))
    for item in TRAINING:
        story.append(Paragraph(f"{item['name']} | {item['date']}", styles["role"]))
        story.append(bullet_list(item["bullets"], styles["body"]))

    story.append(PageBreak())
    story.append(Paragraph("SELECTED PROJECTS", styles["section"]))
    for project in PROJECTS:
        story.append(Paragraph(project["name"], styles["role"]))
        story.append(Paragraph(project["stack"], styles["org"]))
        story.append(bullet_list(project["bullets"], styles["body"]))

    story.append(Paragraph("CERTIFICATIONS", styles["section"]))
    story.append(bullet_list(CERTIFICATIONS, styles["body"]))
    story.append(Paragraph("EDUCATION", styles["section"]))
    for line in EDUCATION:
        story.append(Paragraph(line, styles["body"]))

    doc.build(story, onLaterPages=draw_later_page_header)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
